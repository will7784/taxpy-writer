"""
Exportador de contenido a archivos descargables.
Soporta Markdown y DOCX.
"""

from __future__ import annotations

import io
import re
from datetime import datetime


def to_markdown(content: str, title: str) -> bytes:
    """Genera bytes UTF-8 de un archivo Markdown."""
    header = f"# {title}\n\n"
    header += f"_Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}_\n\n"
    header += "---\n\n"
    full = header + content + "\n"
    return full.encode("utf-8")


def to_docx(content: str, title: str) -> bytes:
    """Genera bytes de un archivo DOCX desde contenido Markdown-like."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Título
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fecha
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run.italic = True
    run.font.size = Pt(10)

    doc.add_paragraph()  # espacio

    # Parsear contenido línea por línea
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Heading 1
        if stripped.startswith("# ") and not stripped.startswith("## "):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # Heading 2
        if stripped.startswith("## ") and not stripped.startswith("### "):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue

        # Heading 3
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue

        # Listas con guion o asterisco
        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        # Listas numeradas
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            p.add_run(re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # Párrafo normal (con soporte básico de **bold**)
        p = doc.add_paragraph()
        _add_formatted_text(p, stripped)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _add_formatted_text(paragraph, text: str) -> None:
    """Agrega texto a un párrafo respetando **bold** y *italic*."""
    # Patrón para **bold** y *italic* (no doble)
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*)"
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
